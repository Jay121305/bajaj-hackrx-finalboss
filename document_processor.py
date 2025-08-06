import os
import tempfile
from typing import Dict, Any, Tuple, List, Optional
import logging
import time
import hashlib
from fastapi import UploadFile, HTTPException
import fitz  # PyMuPDF for PDF processing
import docx  # python-docx for DOCX processing
import email
import re
import requests
from email.parser import BytesParser
from email.policy import default

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process different document types and prepare for embedding"""

    def __init__(self):
        self.document_stats = {
            "processed_count": 0,
            "pdf_count": 0,
            "docx_count": 0,
            "email_count": 0,
            "other_count": 0,
        }
        self.domain_counts = {}
        self.document_store = {}  # Simple in-memory store
    
    async def process_document(self, file: UploadFile, file_extension: str) -> Tuple[str, Dict[str, Any]]:
        """Process document based on its type and return extracted text with metadata"""
        try:
            # Create temp file to process
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}') as temp_file:
                contents = await file.read()
                temp_file.write(contents)
                temp_path = temp_file.name

            # Process based on file type
            if file_extension.lower() == 'pdf':
                text, metadata = self._process_pdf(temp_path)
                self.document_stats["pdf_count"] += 1
            elif file_extension.lower() in ['docx', 'doc']:
                text, metadata = self._process_docx(temp_path)
                self.document_stats["docx_count"] += 1
            elif file_extension.lower() in ['eml', 'msg']:
                text, metadata = self._process_email(temp_path)
                self.document_stats["email_count"] += 1
            else:
                # Try to read as text
                with open(temp_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                metadata = {
                    "file_type": file_extension,
                    "file_size": len(contents),
                    "processing_method": "plain_text"
                }
                self.document_stats["other_count"] += 1

            # Clean up temp file
            os.unlink(temp_path)
            
            # Update document stats
            self.document_stats["processed_count"] += 1
            
            # Store document hash for deduplication checks
            metadata["document_hash"] = hashlib.md5(text.encode()).hexdigest()
            
            return text, metadata

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Document processing error: {str(e)}")
            
    async def process_document_from_url(self, document_url: str) -> Tuple[str, Dict[str, Any]]:
        """Process a document from a URL (for hackathon endpoint)"""
        try:
            # Get file extension from URL
            file_extension = document_url.split('.')[-1].lower()
            
            # Download the document
            response = requests.get(document_url, timeout=30)
            response.raise_for_status()
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}') as temp_file:
                temp_file.write(response.content)
                temp_path = temp_file.name
                
            # Process based on file type
            if file_extension.lower() == 'pdf':
                text, metadata = self._process_pdf(temp_path)
                self.document_stats["pdf_count"] += 1
            elif file_extension.lower() in ['docx', 'doc']:
                text, metadata = self._process_docx(temp_path)
                self.document_stats["docx_count"] += 1
            else:
                # Default to PDF processing for hackathon
                text, metadata = self._process_pdf(temp_path)
                self.document_stats["pdf_count"] += 1
                
            # Clean up temp file
            os.unlink(temp_path)
            
            # Update document stats
            self.document_stats["processed_count"] += 1
            
            # Add URL to metadata
            metadata["source_url"] = document_url
            metadata["document_hash"] = hashlib.md5(text.encode()).hexdigest()
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error processing document from URL: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Document processing error: {str(e)}")

    def _process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process PDF document and extract text with metadata"""
        try:
            start_time = time.time()
            doc = fitz.open(file_path)
            
            # Extract text with structure preservation
            text_blocks = []
            metadata = {
                "page_count": len(doc),
                "file_type": "pdf",
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
                "creation_date": doc.metadata.get("creationDate", ""),
                "modification_date": doc.metadata.get("modDate", ""),
                "has_toc": bool(doc.get_toc()),
                "page_sizes": [],
                "has_images": False,
                "has_tables": False,
                "processing_method": "pymupdf"
            }
            
            # Advanced text extraction with layout analysis
            for page_num, page in enumerate(doc):
                page_size = {"width": page.rect.width, "height": page.rect.height}
                metadata["page_sizes"].append(page_size)
                
                # Check for images
                if not metadata["has_images"] and page.get_images(full=True):
                    metadata["has_images"] = True
                
                # Extract text with structure preservation
                blocks = page.get_text("blocks")
                
                # Check for potential tables
                if not metadata["has_tables"]:
                    # Simple heuristic for table detection - many blocks with similar vertical alignment
                    y_positions = [block[1] for block in blocks if len(block[4].strip()) > 0]
                    y_position_counts = {}
                    for y in y_positions:
                        rounded_y = round(y, 0)
                        y_position_counts[rounded_y] = y_position_counts.get(rounded_y, 0) + 1
                    
                    # If multiple blocks share the same vertical alignment, likely a table
                    if any(count > 2 for count in y_position_counts.values()):
                        metadata["has_tables"] = True
                
                # Add text blocks with page markers
                for block in blocks:
                    if len(block[4].strip()) > 0:
                        text_blocks.append(f"[Page {page_num + 1}] {block[4]}")
            
            full_text = "\n\n".join(text_blocks)
            
            # Detect headings and structure
            headings = self._detect_headings(full_text)
            if headings:
                metadata["detected_headings"] = len(headings)
                metadata["heading_examples"] = headings[:5]  # Store a few examples
            
            # Add processing stats
            metadata["processing_time"] = time.time() - start_time
            metadata["extracted_chars"] = len(full_text)
            
            doc.close()
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            raise HTTPException(status_code=500, detail=f"PDF processing error: {str(e)}")

    def _process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process DOCX document and extract text with metadata"""
        try:
            start_time = time.time()
            doc = docx.Document(file_path)
            
            # Extract document metadata
            metadata = {
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections),
                "has_tables": len(doc.tables) > 0,
                "table_count": len(doc.tables),
                "processing_method": "python-docx"
            }
            
            # Extract document properties if available
            try:
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else ""
                })
            except:
                pass  # Skip if properties not available
            
            # Extract structured content
            content = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    # Check if paragraph is a heading
                    if para.style.name.startswith('Heading'):
                        content.append(f"\n## {para.text.strip()} ##\n")
                    else:
                        content.append(para.text.strip())
            
            # Process tables
            for i, table in enumerate(doc.tables):
                table_text = [f"\n[TABLE {i+1}]"]
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_cells))
                content.append("\n".join(table_text))
            
            full_text = "\n\n".join(content)
            
            # Add processing stats
            metadata["processing_time"] = time.time() - start_time
            metadata["extracted_chars"] = len(full_text)
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise HTTPException(status_code=500, detail=f"DOCX processing error: {str(e)}")

    def _process_email(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Process email file and extract content with metadata"""
        try:
            start_time = time.time()
            
            with open(file_path, 'rb') as fp:
                msg = BytesParser(policy=default).parse(fp)
            
            # Extract email metadata
            metadata = {
                "file_type": "email",
                "subject": msg.get('subject', ''),
                "from": msg.get('from', ''),
                "to": msg.get('to', ''),
                "cc": msg.get('cc', ''),
                "date": msg.get('date', ''),
                "has_attachments": False,
                "processing_method": "email.parser"
            }
            
            # Extract email content
            content_parts = []
            
            # Add header info
            content_parts.append(f"Subject: {metadata['subject']}")
            content_parts.append(f"From: {metadata['from']}")
            content_parts.append(f"To: {metadata['to']}")
            content_parts.append(f"Date: {metadata['date']}")
            content_parts.append("")  # Empty line
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    
                    # Check for attachments
                    if part.get_filename():
                        metadata["has_attachments"] = True
                        continue
                        
                    if content_type == "text/plain":
                        content_parts.append(part.get_content())
                    elif content_type == "text/html":
                        # Simple HTML to text conversion
                        html_content = part.get_content()
                        text_content = self._html_to_text(html_content)
                        content_parts.append(text_content)
            else:
                # Single part email
                content_type = msg.get_content_type()
                if content_type == "text/plain":
                    content_parts.append(msg.get_content())
                elif content_type == "text/html":
                    html_content = msg.get_content()
                    text_content = self._html_to_text(html_content)
                    content_parts.append(text_content)
            
            full_text = "\n\n".join(content_parts)
            
            # Add processing stats
            metadata["processing_time"] = time.time() - start_time
            metadata["extracted_chars"] = len(full_text)
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error processing email: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Email processing error: {str(e)}")

    def _html_to_text(self, html_content: str) -> str:
        """Simple conversion of HTML to plain text"""
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', ' ', html_content)
        # Fix spacing
        text = re.sub(r'\s+', ' ', text)
        # Handle common HTML entities
        text = text.replace('&nbsp;', ' ')
        text = text.replace('&lt;', '<')
        text = text.replace('&gt;', '>')
        text = text.replace('&amp;', '&')
        text = text.replace('&quot;', '"')
        return text.strip()

    def _detect_headings(self, text: str) -> List[str]:
        """Detect potential headings in document text"""
        heading_patterns = [
            r'^[A-Z][A-Z\s]{5,}[A-Z]:?$',  # ALL CAPS HEADING:
            r'^\d+\.\s+[A-Z][a-zA-Z\s]+$',  # 1. Numbered Heading
            r'^[A-Z][a-zA-Z\s]+:$',  # Heading With Colon:
            r'^SECTION\s+\d+',  # SECTION 1
            r'^Article\s+\d+',  # Article 1
            r'^CLAUSE\s+\d+',  # CLAUSE 1
        ]
        
        headings = []
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            for pattern in heading_patterns:
                if re.match(pattern, line, re.MULTILINE):
                    headings.append(line)
                    break
        
        return headings

    async def detect_domain(self, text: str) -> str:
        """Detect document domain based on content analysis"""
        # Domain-specific keywords and patterns
        domain_indicators = {
            "insurance": [
                "policy", "premium", "coverage", "claim", "insured", "underwriting", 
                "deductible", "beneficiary", "rider", "exclusion", "grace period",
                "waiting period", "pre-existing", "maternity", "hospitalization"
            ],
            "legal": [
                "agreement", "contract", "clause", "party", "term", "provision", 
                "covenant", "statute", "jurisdiction", "liability", "plaintiff", "defendant",
                "binding", "arbitration", "warranty", "indemnification", "termination"
            ],
            "hr": [
                "employee", "employment", "compensation", "benefits", "payroll", 
                "recruitment", "termination", "hr policy", "leave", "workplace",
                "attendance", "performance", "appraisal", "vacation", "salary"
            ],
            "compliance": [
                "regulation", "compliance", "audit", "standard", "requirement", 
                "procedure", "guideline", "framework", "certification", "governance",
                "risk", "violation", "penalty", "reporting", "disclosure"
            ],
            "finance": [
                "investment", "financial", "asset", "portfolio", "capital", 
                "dividend", "stock", "equity", "bond", "security",
                "maturity", "interest", "principal", "yield", "amortization"
            ],
            "medical": [
                "patient", "treatment", "diagnosis", "physician", "hospital", 
                "clinic", "prescription", "symptom", "prognosis", "healthcare",
                "disease", "procedure", "medication", "dosage", "discharge"
            ]
        }
        
        # Count occurrences of domain-specific keywords
        domain_scores = {domain: 0 for domain in domain_indicators}
        text_lower = text.lower()
        
        for domain, keywords in domain_indicators.items():
            for keyword in keywords:
                domain_scores[domain] += text_lower.count(keyword.lower())
        
        # Get domain with highest score
        max_score = 0
        detected_domain = "general"  # Default
        
        for domain, score in domain_scores.items():
            if score > max_score:
                max_score = score
                detected_domain = domain
        
        # Update domain statistics
        self.domain_counts[detected_domain] = self.domain_counts.get(detected_domain, 0) + 1
        
        return detected_domain

    async def get_document_count(self) -> int:
        """Get total count of processed documents"""
        return self.document_stats["processed_count"]

    async def get_available_domains(self) -> List[str]:
        """Get list of detected domains in the system"""
        return list(self.domain_counts.keys())